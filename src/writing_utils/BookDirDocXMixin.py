import os
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from utils import File, Log

from private import data

log = Log("BookDirDocXMixin")


class BookDirDocXMixin:

    def build_docx(self, max_words_per_docx: int = 50000) -> list[Document]:
        return self.__create_docx_documents_and_save__(max_words_per_docx)

    def __create_docx_documents_and_save__(
        self, max_words_per_docx: int = 50000
    ) -> list[Document]:
        compiled_dir = self.path + ".compiled"
        docx_dir = os.path.join(compiled_dir, "docx")
        shutil.rmtree(docx_dir, ignore_errors=True)
        os.makedirs(docx_dir, exist_ok=True)

        self.__create_and_save_docx_files__(docx_dir, max_words_per_docx)
        return docx_dir

    def __create_and_save_docx_files__(
        self, docx_dir: str, max_words_per_docx: int = 50000
    ) -> list[str]:
        chapters = sorted(self.gen_chapter_docs(), key=lambda ch: ch.number)
        docx_paths = []
        current_doc = None
        current_docx_index = 0
        current_word_count = 0

        for chapter_doc in chapters:
            chapter_word_count = chapter_doc.n_words

            if (
                current_doc is not None
                and current_word_count + chapter_word_count
                > max_words_per_docx
            ):
                docx_path = os.path.join(
                    docx_dir, f"part_{current_docx_index:02d}.docx"
                )
                current_doc.save(docx_path)
                docx_paths.append(docx_path)
                log.info(
                    f"📄 Wrote {File(docx_path)} ({current_word_count} words)"
                )

                current_doc = None
                current_docx_index += 1
                current_word_count = 0

            if current_doc is None:
                current_doc = self.__create_docx_document__()

            self.__add_docx_chapter_section__(current_doc, chapter_doc)
            current_word_count += chapter_word_count

        if current_doc is not None:
            docx_path = os.path.join(
                docx_dir, f"part_{current_docx_index:02d}.docx"
            )
            current_doc.save(docx_path)
            docx_paths.append(docx_path)
            log.debug(
                f"📄 Wrote {File(docx_path)} ({current_word_count} words)"
            )

        log.info(f"📚 Wrote {len(docx_paths)} parts to {docx_dir}")
        return docx_paths

    def __create_docx_document__(self) -> Document:
        doc = Document()
        self.__configure_docx_page_layout__(doc)
        self.__add_docx_title_page__(doc)
        return doc

    def __configure_docx_page_layout__(self, doc: Document):
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(12)

    def __add_docx_title_page__(self, doc: Document):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(data.TITLE)
        title_run.font.size = Pt(28)
        title_run.font.bold = True

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(data.SUBTITLE)
        subtitle_run.font.size = Pt(18)

        doc.add_paragraph()
        doc.add_paragraph()

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"By {data.AUTHOR}")
        author_run.font.size = Pt(14)

        for _ in range(5):
            doc.add_paragraph()

        copyright_para = doc.add_paragraph()
        copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_run = copyright_para.add_run(
            f"Copyright © 2025 by {data.AUTHOR}"
        )
        copyright_run.font.size = Pt(10)

        copyright_notice = doc.add_paragraph()
        copyright_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_notice_run = copyright_notice.add_run(
            "All rights reserved."
        )
        copyright_notice_run.font.size = Pt(10)

        doc.add_page_break()

    def __add_docx_chapter_section__(self, doc: Document, chapter_doc):
        chapter_heading = f"{chapter_doc.number}. {chapter_doc.title}"
        doc.add_heading(chapter_heading, level=1)
        lines = chapter_doc.lines[1:]
        content = "\n".join(lines).strip()
        self.__convert_markdown_to_docx__(content, doc)

    def __convert_markdown_to_docx__(
        self, content: str, doc: Document
    ) -> None:
        paragraphs = content.split("\n")

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            if re.match(r"^-{3,}$", para_text):
                doc.add_paragraph("---")
                continue

            header_match = re.match(r"^(#{1,3})\s+(.+)$", para_text)
            if header_match:
                level = len(header_match.group(1))
                title = header_match.group(2)
                doc.add_heading(title, level=level)
                continue

            para = doc.add_paragraph()
            self.__add_formatted_text_to_docx_paragraph__(para, para_text)

    def __add_formatted_text_to_docx_paragraph__(
        self, para, text: str
    ) -> None:
        pos = 0

        while pos < len(text):
            bold_match = re.search(r"\*\*(.+?)\*\*", text[pos:])
            italic_match = re.search(r"\*(.+?)\*", text[pos:])
            say_match = re.search(r"\\say\{(.+?)\}", text[pos:])

            matches = []
            if bold_match:
                matches.append(
                    (
                        "bold",
                        bold_match.start(),
                        bold_match.end(),
                        bold_match.group(1),
                    )
                )
            if italic_match and not (
                bold_match
                and italic_match.start() >= bold_match.start()
                and italic_match.start() < bold_match.end()
            ):
                matches.append(
                    (
                        "italic",
                        italic_match.start(),
                        italic_match.end(),
                        italic_match.group(1),
                    )
                )
            if say_match:
                matches.append(
                    (
                        "say",
                        say_match.start(),
                        say_match.end(),
                        say_match.group(1),
                    )
                )

            if not matches:
                para.add_run(text[pos:])
                break

            matches.sort(key=lambda x: x[1])
            match_type, start, end, content = matches[0]
            abs_start = pos + start
            abs_end = pos + end

            if start > 0:
                para.add_run(text[pos:abs_start])

            run = para.add_run(content)
            if match_type == "bold":
                run.bold = True
            elif match_type == "italic":
                run.italic = True
            elif match_type == "say":
                run.italic = True

            pos = abs_end
