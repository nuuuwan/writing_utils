import glob
import os
import re
import shutil

from docx import Document
from utils import File, Log

from writing_utils.ChapterFile import ChapterFile

log = Log("BookDirReverseDocXMixin")


class BookDirReverseDocXMixin:
    @staticmethod
    def extract_formatted_text(para):
        formatted_text = ""
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            formatted_text += text
        return formatted_text.strip()

    @classmethod
    def from_docx(cls, docx_path: str):
        """Load BookDir from a single DOCX file or directory of multiple DOCX files."""
        # Determine if it's a directory or single file
        if os.path.isdir(docx_path):
            # Load from directory of multiple DOCX files
            return cls._load_from_docx_directory(docx_path)
        else:
            # Load from single DOCX file (legacy support)
            return cls._load_from_single_docx(docx_path)

    @classmethod
    def _load_from_single_docx(cls, docx_path: str):
        """Load BookDir from a single DOCX file."""
        doc = Document(docx_path)
        temp_dir = docx_path + ".bookdir"
        shutil.rmtree(temp_dir, ignore_errors=True)
        os.makedirs(temp_dir, exist_ok=True)

        chapters = cls._parse_chapters_from_doc(doc)
        cls._save_chapters_to_files(chapters, temp_dir)

        book_dir = cls(temp_dir)
        log.info(f"📖 Loaded BookDir from {docx_path}")
        return book_dir

    @classmethod
    def _load_from_docx_directory(cls, docx_dir: str):
        """Load BookDir from a directory containing multiple DOCX files."""
        # Find all part_*.docx files in the directory, sorted by part number
        docx_files = sorted(glob.glob(os.path.join(docx_dir, "part_*.docx")))

        if not docx_files:
            raise ValueError(f"No DOCX files found in {docx_dir}")

        temp_dir = docx_dir + ".bookdir"
        shutil.rmtree(temp_dir, ignore_errors=True)
        os.makedirs(temp_dir, exist_ok=True)

        all_chapters = {}

        for docx_path in docx_files:
            doc = Document(docx_path)
            chapters = cls._parse_chapters_from_doc(doc, skip_title_page=True)
            all_chapters.update(chapters)

        cls._save_chapters_to_files(all_chapters, temp_dir)

        book_dir = cls(temp_dir)
        log.info(
            f"📖 Loaded BookDir from {len(docx_files)} DOCX files in {docx_dir}"
        )
        return book_dir

    @classmethod
    def _parse_chapters_from_doc(cls, doc, skip_title_page=False):
        """Extract chapters from Document, preserving formatting.

        Args:
            doc: The Document to parse
            skip_title_page: If True, skip paragraphs before the first chapter heading
        """
        chapters = {}
        current_chapter_num = None
        current_chapter_title = None
        current_chapter_lines = []
        found_first_chapter = (
            not skip_title_page
        )  # If skipping, start in "found" state

        for para in doc.paragraphs:
            style_name = para.style.name
            text = cls.extract_formatted_text(para)

            if style_name.startswith("Heading 1"):
                heading_match = re.match(r"(\d+)\.\s+(.+)", text)
                if heading_match:
                    found_first_chapter = True
                    if current_chapter_num is not None:
                        chapters[current_chapter_num] = (
                            current_chapter_title,
                            "\n".join(current_chapter_lines),
                        )
                    current_chapter_num = int(heading_match.group(1))
                    current_chapter_title = heading_match.group(2)
                    current_chapter_lines = [
                        f"# {current_chapter_num}. {current_chapter_title}"
                    ]
            elif (
                found_first_chapter
                and current_chapter_num is not None
                and text
            ):
                current_chapter_lines.append(text)

        if current_chapter_num is not None:
            chapters[current_chapter_num] = (
                current_chapter_title,
                "\n".join(current_chapter_lines),
            )

        return chapters

    @staticmethod
    def _save_chapters_to_files(chapters, temp_dir):
        """Write chapters to individual markdown files."""
        for chapter_num, (title, content) in sorted(chapters.items()):
            filename = f"{chapter_num:02d}-{title.replace(' ', '_')}.md"
            filepath = os.path.join(temp_dir, filename)
            cleaned_content = ChapterFile.get_cleaned(content)
            File(filepath).write(cleaned_content)
